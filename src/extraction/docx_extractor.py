from docx import Document
from pathlib import Path
from typing import List
import uuid
from src.utils.schema import DocumentChunk

def extract_docx_chunks(docx_path: str, doc_id: str) -> List[DocumentChunk]:
    doc = Document(docx_path)
    chunks = []
    source_file = Path(docx_path).name
    current_heading = None
    page_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 20:
            continue
        if para.style.name.startswith("Heading"):
            current_heading = text
        chunks.append(DocumentChunk(
            chunk_id=str(uuid.uuid4()),
            doc_id=doc_id,
            source_file=source_file,
            page_number=page_num,
            modality="text",
            content=text,
            section_heading=current_heading,
            extraction_confidence=1.0,
        ))

    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        text_repr = " | ".join(headers) + "\n"
        for row in rows:
            text_repr += " | ".join(row) + "\n"
        chunks.append(DocumentChunk(
            chunk_id=str(uuid.uuid4()),
            doc_id=doc_id,
            source_file=source_file,
            page_number=page_num,
            modality="table",
            content=text_repr.strip(),
            structured_data={"headers": headers, "rows": rows},
            extraction_confidence=0.9,
        ))

    return chunks
